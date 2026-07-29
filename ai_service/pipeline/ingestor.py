"""
pipeline/ingestor.py
--------------------
Handles all document ingestion for ContrAIct:
  PDF / DOCX / scanned images (OCR) → text extraction
  → chunking → embeddings → per-contract Chroma vector store

Each contract gets its own isolated vector store directory so
retrieval is always scoped to a single document.
"""
import os
import shutil
import time
import json
from pathlib import Path
from typing import List, Tuple

from langchain_community.document_loaders import PyPDFLoader
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

import config


# ── Embeddings singleton (loaded once, shared across all calls) ────────────────
_embeddings: HuggingFaceEmbeddings | None = None

def get_embeddings() -> HuggingFaceEmbeddings:
    global _embeddings
    if _embeddings is None:
        _embeddings = HuggingFaceEmbeddings(model_name=config.EMBEDDING_MODEL)
    return _embeddings


# ── Text extraction ────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> Tuple[str, int]:
    """Extract raw text from PDF. Returns (full_text, page_count)."""
    loader = PyPDFLoader(file_path)
    pages  = loader.load()
    full_text = "\n\n".join(p.page_content for p in pages)
    if full_text.strip():
        return full_text, len(pages)

    print("[OCR FALLBACK] PDF text extraction returned no text; trying OCR")
    ocr_text = extract_text_from_pdf_ocr(file_path)
    if ocr_text.strip():
        return ocr_text, len(pages)

    return full_text, len(pages)


def extract_text_from_pdf_ocr(file_path: str) -> str:
    """OCR a scanned PDF when normal text extraction fails."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except ImportError:
        return ""

    images = convert_from_path(file_path)
    return "\n\n".join(pytesseract.image_to_string(image) for image in images)


def extract_text_from_docx(file_path: str) -> Tuple[str, int]:
    """Extract raw text from DOCX. Returns (full_text, page_count_estimate)."""
    from docx import Document as DocxDocument
    doc   = DocxDocument(file_path)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    text  = "\n\n".join(paras)
    # DOCX has no real page count — estimate from word count
    words      = len(text.split())
    page_count = max(1, words // 300)
    return text, page_count


def extract_text_from_image(file_path: str) -> Tuple[str, int]:
    """OCR a scanned image using Tesseract. Returns (text, 1)."""
    try:
        import pytesseract
        from PIL import Image
        img  = Image.open(file_path)
        text = pytesseract.image_to_string(img)
        return text, 1
    except ImportError:
        raise RuntimeError(
            "pytesseract or Pillow not installed. "
            "Run: pip install pytesseract pillow"
        )


def extract_text(file_path: str, mime_type: str) -> Tuple[str, int]:
    """
    Route to the correct extractor based on mime type.
    Returns (raw_text, page_count).
    """
    ext = Path(file_path).suffix.lower()

    if mime_type == "application/pdf" or ext == ".pdf":
        return extract_text_from_pdf(file_path)

    if (
        mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or ext == ".docx"
    ):
        return extract_text_from_docx(file_path)

    if mime_type in ("image/png", "image/jpeg", "image/tiff") or ext in (
        ".png", ".jpg", ".jpeg", ".tiff"
    ):
        return extract_text_from_image(file_path)

    raise ValueError(f"Unsupported file type: {mime_type} / {ext}")


# ── Chunking ──────────────────────────────────────────────────────────────────

def chunk_text(raw_text: str, contract_id: str) -> List[Document]:
    """
    Split raw text into overlapping chunks.
    Each chunk carries contract_id in metadata for scoped retrieval.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    # Wrap raw text in a Document first so metadata flows through
    doc    = Document(page_content=raw_text, metadata={"contract_id": contract_id})
    chunks = splitter.split_documents([doc])

    # Ensure every chunk carries the contract_id
    for i, chunk in enumerate(chunks):
        chunk.metadata["contract_id"] = contract_id
        chunk.metadata["chunk_index"] = i

    return chunks


# ── Vector store ──────────────────────────────────────────────────────────────

def vectorstore_path(contract_id: str) -> str:
    """Each contract gets its own Chroma collection directory."""
    return os.path.join(config.VECTORSTORE_BASE_DIR, contract_id)


def cache_dir(document_hash: str | None) -> str | None:
    if not document_hash:
        return None
    return os.path.join(config.CACHE_BASE_DIR, document_hash)


def cached_text_path(document_hash: str | None) -> str | None:
    directory = cache_dir(document_hash)
    if not directory:
        return None
    return os.path.join(directory, "extracted_text.json")


def cached_vectorstore_path(document_hash: str | None) -> str | None:
    directory = cache_dir(document_hash)
    if not directory:
        return None
    return os.path.join(directory, "vectorstore")


def load_cached_text(document_hash: str | None) -> Tuple[str, int] | None:
    path = cached_text_path(document_hash)
    if not path or not os.path.exists(path):
        return None
    with open(path, "r", encoding="utf-8") as cache_file:
        payload = json.load(cache_file)
    return payload["rawText"], int(payload.get("pageCount", 0))


def save_cached_text(document_hash: str | None, raw_text: str, page_count: int) -> None:
    path = cached_text_path(document_hash)
    if not path:
        return
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as cache_file:
        json.dump({"rawText": raw_text, "pageCount": page_count}, cache_file)


def build_vectorstore(chunks: List[Document], contract_id: str, document_hash: str | None = None) -> Chroma:
    """Embed chunks and persist to disk under vectorstores/<contract_id>/."""
    vs_dir = vectorstore_path(contract_id)
    cached_vs_dir = cached_vectorstore_path(document_hash)

    # Wipe existing store for this contract (re-ingest scenario)
    if os.path.exists(vs_dir):
        shutil.rmtree(vs_dir)

    if cached_vs_dir and os.path.exists(cached_vs_dir):
        shutil.copytree(cached_vs_dir, vs_dir)
        print(f"[VECTORSTORE CACHE HIT] contractId={contract_id} documentHash={document_hash}")
        return Chroma(
            persist_directory=vs_dir,
            embedding_function=get_embeddings(),
        )

    vectorstore = Chroma.from_documents(
        documents=chunks,
        embedding=get_embeddings(),
        persist_directory=vs_dir,
    )
    if cached_vs_dir:
        if os.path.exists(cached_vs_dir):
            shutil.rmtree(cached_vs_dir)
        shutil.copytree(vs_dir, cached_vs_dir)
    return vectorstore


def load_vectorstore(contract_id: str) -> Chroma:
    """Load an existing vector store for a contract."""
    vs_dir = vectorstore_path(contract_id)
    if not os.path.exists(vs_dir):
        raise FileNotFoundError(
            f"No vector store found for contract {contract_id}. "
            "Ingest the document first."
        )
    return Chroma(
        persist_directory=vs_dir,
        embedding_function=get_embeddings(),
    )


# ── Full ingestion pipeline ───────────────────────────────────────────────────

def ingest_contract(
    file_path: str,
    mime_type: str,
    contract_id: str,
    document_hash: str | None = None,
) -> Tuple[str, int, Chroma]:
    """
    Full pipeline for a single contract:
      extract text → chunk → embed → store

    Returns:
        raw_text    : full extracted text (stored in DB for re-analysis)
        page_count  : number of pages
        vectorstore : built Chroma instance ready for retrieval
    """
    started_at = time.perf_counter()

    cached_text = load_cached_text(document_hash)
    if cached_text:
        raw_text, page_count = cached_text
        print(f"[TEXT CACHE HIT] contractId={contract_id} documentHash={document_hash}")
    else:
        extraction_started_at = time.perf_counter()
        raw_text, page_count = extract_text(file_path, mime_type)
        print(f"[TEXT EXTRACTION] contractId={contract_id} elapsed={time.perf_counter() - extraction_started_at:.2f}s")
        save_cached_text(document_hash, raw_text, page_count)

    if not raw_text.strip():
        raise ValueError("No text could be extracted from the document.")

    chunking_started_at = time.perf_counter()
    chunks      = chunk_text(raw_text, contract_id)
    print(f"[CHUNKING] contractId={contract_id} elapsed={time.perf_counter() - chunking_started_at:.2f}s chunks={len(chunks)}")

    embeddings_started_at = time.perf_counter()
    vectorstore = build_vectorstore(chunks, contract_id, document_hash)
    print(f"[EMBEDDINGS] contractId={contract_id} elapsed={time.perf_counter() - embeddings_started_at:.2f}s")
    print(f"[INGEST DONE] contractId={contract_id} elapsed={time.perf_counter() - started_at:.2f}s")

    return raw_text, page_count, vectorstore
